#!/usr/bin/env python3
"""解析 docx 中的所有表格结构，输出合并格字段和单元格内容"""
import sys
from docx import Document
from docx.oxml.ns import qn

def get_cell_spans(tc):
    tcPr = tc.find(qn('w:tcPr'))
    grid_span_el = tcPr.find(qn('w:gridSpan')) if tcPr is not None else None
    colspan = int(grid_span_el.get(qn('w:val'), 1)) if grid_span_el is not None else 1

    vmerge = tcPr.find(qn('w:vMerge')) if tcPr is not None else None
    if vmerge is not None:
        val = vmerge.get(qn('w:val'), '')
        rowspan = 'START' if val == 'restart' else 'CONT'
    else:
        rowspan = 1

    return rowspan, colspan

def cell_text(tc):
    return ''.join(
        ''.join(r.text for r in p.findall('.//' + qn('w:r')))
        for p in tc.findall('.//' + qn('w:p'))
    ).strip()

def print_table(idx, table):
    print(f"\n{'='*70}")
    print(f"  TABLE {idx+1}  ({len(table.rows)} 行 x {len(table.columns)} 列)")
    print('='*70)

    row_count = len(table.rows)
    # map (r, logical_col) -> vmerge start row
    vmerge_starts = {}

    for r, row in enumerate(table.rows):
        c_idx = 0
        cells_info = []
        for tc in row._tr.findall(qn('w:tc')):
            rowspan, colspan = get_cell_spans(tc)
            text = cell_text(tc)

            if rowspan == 'START':
                for cc in range(c_idx, c_idx + colspan):
                    vmerge_starts[(r, cc)] = r
                # calc actual span
                count = 1
                for rr in range(r+1, row_count):
                    found = False
                    for cc in range(c_idx, c_idx + colspan):
                        if vmerge_starts.get((rr, cc)) == r:
                            found = True
                    if found:
                        count += 1
                    else:
                        break
                rs_str = f"rowspan={count}"
            elif rowspan == 'CONT':
                for cc in range(c_idx, c_idx + colspan):
                    # propagate
                    for rr in range(r-1, -1, -1):
                        if (rr, cc) in vmerge_starts:
                            vmerge_starts[(r, cc)] = vmerge_starts[(rr, cc)]
                            break
                rs_str = "(merged↑)"
            else:
                rs_str = ""

            cs_str = f"colspan={colspan}" if colspan > 1 else ""
            meta = " ".join(filter(None, [rs_str, cs_str]))
            cells_info.append(f"  col{c_idx}[{meta}]: 「{text}」" if meta else f"  col{c_idx}: 「{text}」")
            c_idx += colspan

        print(f"行{r}:")
        for ci in cells_info:
            print(ci)

def main():
    path = r"e:\code\antigravity_projects\ReportTemplate\output\key-device-report-sample.docx"
    doc = Document(path)
    print(f"共发现 {len(doc.tables)} 个表格")

    print("\n--- 段落标题 ---")
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') or para.text.strip():
            print(f"  [{para.style.name}] {para.text.strip()[:80]}")

    for i, table in enumerate(doc.tables):
        print_table(i, table)

if __name__ == '__main__':
    main()
